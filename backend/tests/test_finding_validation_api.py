"""Live integration tests for Phase 6's FindingValidationAgent — structural
copy of `test_query_api.py`'s pattern (real Postgres+Qdrant, paid third-party
calls mocked) combined with `test_iso27001_findings_api.py`'s scan-upload
helper, since this phase needs both a real ingested standard (for retrieval)
and a real scan (for findings to validate).
"""

import io
import uuid
import zipfile

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import text

from app.core.db import engine
from app.main import app
from app.services import embedding, finding_validation, reranking, retrieval, storage, vector_store
from app.services.finding_validation import (
    ContextRelationship,
    FindingAssessment,
    FindingValidationVerdict,
)
from app.tasks.celery_app import celery_app


def _infra_available() -> bool:
    try:
        with engine.connect() as conn:
            conn.execute(text("SELECT 1 FROM findings LIMIT 1"))
        storage.get_minio_client().list_buckets()
        vector_store.get_qdrant_client().get_collections()
        return True
    except Exception:
        return False


pytestmark = pytest.mark.skipif(
    not _infra_available(),
    reason="requires docker compose up (postgres+redis+minio+qdrant) with migrations applied",
)


@pytest.fixture(autouse=True, scope="module")
def _eager_celery():
    celery_app.conf.task_always_eager = True
    celery_app.conf.task_eager_propagates = True
    yield
    celery_app.conf.task_always_eager = False
    celery_app.conf.task_eager_propagates = False


@pytest.fixture(autouse=True)
def _mock_external_apis(monkeypatch):
    """Mocks every paid third-party call (Voyage, Cohere, Gemini) so this
    test needs no real API keys, while still exercising the real Postgres
    lexical search + real Qdrant dense search + real RRF fusion wiring —
    same philosophy as `test_query_api.py`'s own fixture of the same name.
    """

    def _fake_embed_texts(texts, input_type):
        return [[0.1] * embedding.EMBEDDING_DIM for _ in texts]

    def _fake_rerank(query, documents, top_n):
        return [
            reranking.RerankResult(index=i, relevance_score=1.0 - i * 0.01)
            for i in range(min(top_n, len(documents)))
        ]

    monkeypatch.setattr(embedding, "embed_texts", _fake_embed_texts)
    monkeypatch.setattr(reranking, "rerank", _fake_rerank)

    class _FakeValidationClient:
        def validate(self, prompt: str) -> FindingValidationVerdict:
            return FindingValidationVerdict(
                context_relationship=ContextRelationship.SUPPORTS_CONCERN,
                finding_assessment=FindingAssessment.LIKELY_TRUE_POSITIVE,
                confidence="high",
                rationale=(
                    "The retrieved excerpt on cryptographic key management supports this concern."
                ),
            )

    monkeypatch.setattr(
        finding_validation, "get_gemini_validation_client", lambda: _FakeValidationClient()
    )


def _sample_docx_bytes(nonce: str) -> bytes:
    """`nonce` gets woven into the clause text itself — the shared dev/test
    Qdrant collection has accumulated real standards content across this
    project's history (including a genuine ISO 27001 document ingested in
    much earlier phases, which has its own real, richer A.8.24 text), so a
    generic "cryptography/key management" paragraph risks losing a fake,
    tie-broken rerank to that older, more keyword-dense chunk. A random
    per-test nonce makes lexical search unambiguously prefer *this* chunk,
    without relying on resetting or otherwise depending on the shared
    corpus's exact contents (a known, deliberately-not-"fixed" limitation
    documented elsewhere in this project).
    """
    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_heading("Validation Test Standard", level=1)
    doc.add_heading("A.8.24 Use of cryptography", level=2)
    doc.add_paragraph(
        f"Rules for the effective use of cryptography {nonce}, including key "
        f"management {nonce}, shall be defined and implemented."
    )
    doc.save(buf)
    return buf.getvalue()


def _sample_repo_zip(marker: str = "") -> bytes:
    """A repo with a weak-crypto finding (SEC-CRYPTO-WEAK-PY, category
    "cryptography") — the exact category the seeded standard above is
    about, so retrieval has something real to find.
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as archive:
        archive.writestr(
            "app/auth.py",
            "import hashlib\n"
            "\n"
            "def hash_password(password):\n"
            "    return hashlib.md5(password.encode()).hexdigest()\n",
        )
        archive.writestr("requirements.txt", "fastapi>=0.100.0\n")
        if marker:
            archive.writestr("MARKER.txt", marker)
    return buf.getvalue()


def _upload_and_wait(client: TestClient, data: bytes, filename: str) -> str:
    response = client.post("/scans", files={"file": (filename, data, "application/zip")})
    scan_id = response.json()["id"]
    status_response = client.get(f"/scans/{scan_id}").json()
    assert status_response["findings_status"] == "ready"
    return scan_id


def _first_security_finding_id(client: TestClient, scan_id: str) -> str:
    all_findings = client.get(f"/scans/{scan_id}/findings").json()
    findings = [
        f for f in all_findings if f["framework"] is None and f["category"] == "cryptography"
    ]
    return findings[0]["id"]


def test_validate_finding_creates_llm_reasoning_evidence_with_real_retrieval():
    """Asserts the endpoint's real plumbing end-to-end (real dense+lexical+
    fusion+rerank retrieval feeds real citations into a real persisted
    Evidence row) — deliberately does NOT assert that this test's own
    ingested document specifically wins the retrieval race. Dense-search
    ranking is meaningless here (every embed call returns the same fake
    constant vector, same caveat `test_query_api.py`'s own tests document),
    and the shared dev/test Qdrant collection already holds a real ISO
    27001 document's chunks from much earlier phases, which may
    legitimately out-rank a fabricated one-paragraph stub on a generic
    term like "cryptography" via real lexical search. Ingesting a document
    here still matters — it guarantees the collection is non-empty so this
    test exercises the real "chunks found" path rather than the 422 path
    `test_validate_finding_422s_when_no_standards_match` covers.
    """
    client = TestClient(app)
    nonce = uuid.uuid4().hex[:12]
    upload = client.post(
        "/documents",
        files={
            "file": (
                f"validation_standard_{nonce}.docx",
                _sample_docx_bytes(nonce),
                "application/octet-stream",
            )
        },
    )
    assert client.get(f"/documents/{upload.json()['id']}").json()["status"] == "ready"

    scan_id = _upload_and_wait(client, _sample_repo_zip(marker="validate-real"), "validate.zip")
    finding_id = _first_security_finding_id(client, scan_id)

    response = client.post(f"/scans/{scan_id}/findings/{finding_id}/validate")

    assert response.status_code == 200
    body = response.json()
    assert body["source_type"] == "llm_reasoning"
    assert body["rule_id"] == "llm_finding_validation"
    assert body["evidence_metadata"]["finding_assessment"] == "likely_true_positive"
    assert body["evidence_metadata"]["context_relationship"] == "supports_concern"
    citations = body["evidence_metadata"]["retrieved_citations"]
    assert len(citations) >= 1
    for citation in citations:
        assert citation["chunk_id"]
        assert citation["document_id"]
        assert citation["document_filename"]

    # The finding's own status must never have changed as a side effect.
    finding_detail = client.get(f"/scans/{scan_id}/findings/{finding_id}").json()
    assert finding_detail["status"] in (
        "POTENTIAL_NON_COMPLIANCE",
        "REQUIRES_HUMAN_REVIEW",
    )
    llm_evidence = [e for e in finding_detail["evidence"] if e["source_type"] == "llm_reasoning"]
    assert len(llm_evidence) == 1


def test_validate_finding_404s_for_wrong_scan_or_finding_id():
    client = TestClient(app)
    scan_id = _upload_and_wait(client, _sample_repo_zip(marker="404-case"), "404.zip")
    bogus_finding_id = str(uuid.uuid4())

    response = client.post(f"/scans/{scan_id}/findings/{bogus_finding_id}/validate")
    assert response.status_code == 404

    bogus_scan_id = str(uuid.uuid4())
    finding_id = _first_security_finding_id(client, scan_id)
    response = client.post(f"/scans/{bogus_scan_id}/findings/{finding_id}/validate")
    assert response.status_code == 404


def test_validate_finding_422s_when_no_standards_match(monkeypatch):
    """Exercises NoStandardsContextError deterministically by making
    retrieval return zero chunks, rather than relying on the shared dev/
    test database genuinely having no standards ingested (an unreliable
    thing to assert given this project's own documented test-pollution
    history).
    """
    monkeypatch.setattr(retrieval, "vector_search", lambda db, question, top_k: ([], [0.1]))

    client = TestClient(app)
    scan_id = _upload_and_wait(client, _sample_repo_zip(marker="422-case"), "422.zip")
    finding_id = _first_security_finding_id(client, scan_id)

    response = client.post(f"/scans/{scan_id}/findings/{finding_id}/validate")

    assert response.status_code == 422


def test_bulk_validate_rejects_over_cap():
    client = TestClient(app)
    scan_id = _upload_and_wait(client, _sample_repo_zip(marker="bulk-cap"), "bulk-cap.zip")

    response = client.post(
        f"/scans/{scan_id}/findings/validate-bulk",
        json={"finding_ids": [str(uuid.uuid4()) for _ in range(11)]},
    )

    assert response.status_code == 400


def test_bulk_validate_partial_success_on_bad_id():
    client = TestClient(app)
    upload = client.post(
        "/documents",
        files={
            "file": (
                "bulk_standard.docx",
                _sample_docx_bytes(uuid.uuid4().hex[:12]),
                "application/octet-stream",
            )
        },
    )
    assert client.get(f"/documents/{upload.json()['id']}").json()["status"] == "ready"

    scan_id = _upload_and_wait(client, _sample_repo_zip(marker="bulk-partial"), "bulk-partial.zip")
    real_finding_id = _first_security_finding_id(client, scan_id)
    bogus_finding_id = str(uuid.uuid4())

    response = client.post(
        f"/scans/{scan_id}/findings/validate-bulk",
        json={"finding_ids": [real_finding_id, bogus_finding_id]},
    )

    assert response.status_code == 200
    results = {r["finding_id"]: r for r in response.json()}
    assert results[real_finding_id]["ok"] is True
    assert results[real_finding_id]["evidence"]["source_type"] == "llm_reasoning"
    assert results[bogus_finding_id]["ok"] is False
    assert results[bogus_finding_id]["error"] == "finding not found"
