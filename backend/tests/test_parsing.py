from docx import Document as DocxDocument

from app.services.parsing import parse_document


def _build_sample_docx(path) -> None:
    doc = DocxDocument()
    doc.add_heading("ISO 27001 Sample", level=1)
    doc.add_heading("A.8 Asset Management", level=2)
    doc.add_paragraph("This clause covers asset management controls.")
    doc.add_heading("A.8.1 Inventory of assets", level=3)
    doc.add_paragraph("Assets associated with information shall be identified.")
    doc.save(path)


def test_parse_document_preserves_heading_hierarchy(tmp_path):
    docx_path = tmp_path / "sample.docx"
    _build_sample_docx(docx_path)

    root = parse_document(docx_path.read_bytes(), "sample.docx")

    top = root.children[0]
    assert top.title == "ISO 27001 Sample"

    a8 = top.children[0]
    assert a8.clause_number == "A.8"

    a8_1 = a8.children[0]
    assert a8_1.clause_number == "A.8.1"
    assert "identified" in a8_1.text
