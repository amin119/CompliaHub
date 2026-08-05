import io
import uuid

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import text

from app.core.db import engine
from app.main import app
from app.services import embedding, extraction, graph_store
from app.services.ontology import (
    ChunkExtraction,
    EntityType,
    ExtractedEntity,
    ExtractedRelation,
    RelationType,
)
from app.tasks.celery_app import celery_app


def _infra_available() -> bool:
    try:
        with engine.connect() as conn:
            conn.execute(text("SELECT 1 FROM documents LIMIT 1"))
        driver = graph_store.get_neo4j_driver()
        driver.verify_connectivity()
        driver.close()
        return True
    except Exception:
        return False


pytestmark = pytest.mark.skipif(
    not _infra_available(),
    reason="requires docker compose up (postgres+neo4j) with migrations applied",
)


@pytest.fixture(autouse=True, scope="module")
def _eager_celery():
    celery_app.conf.task_always_eager = True
    celery_app.conf.task_eager_propagates = True
    yield
    celery_app.conf.task_always_eager = False
    celery_app.conf.task_eager_propagates = False


@pytest.fixture(autouse=True)
def _mock_external_apis(monkeypatch):
    """Mocks Voyage (embeddings, used by entity resolution) and Anthropic
    (extraction) so this test needs no real API keys, while still
    exercising real Postgres + real Neo4j wiring end-to-end.
    """

    def _fake_embed_texts(texts, input_type):
        return [[0.1] * embedding.EMBEDDING_DIM for _ in texts]

    fixed_extraction = ChunkExtraction(
        entities=[
            ExtractedEntity(name="Test Control", entity_type=EntityType.CONTROL),
            ExtractedEntity(name="Test Risk", entity_type=EntityType.RISK),
        ],
        relations=[
            ExtractedRelation(
                source="Test Control", relation_type=RelationType.APPLIES_TO, target="Test Risk"
            )
        ],
    )

    def _fake_extract_chunk_text(chunk_text, client=None):
        return fixed_extraction

    monkeypatch.setattr(embedding, "embed_texts", _fake_embed_texts)
    monkeypatch.setattr(extraction, "extract_chunk_text", _fake_extract_chunk_text)


def _sample_docx_bytes() -> bytes:
    """A random UUID embedded in the *text content* (not just the filename)
    guarantees a fresh sha256 hash — and therefore a brand new `Document` and
    `Chunk` rows — every test run. Without this, `create_relation`'s
    always-CREATE (no dedup) would accumulate duplicate parallel edges across
    repeated test runs against the same never-cleaned-up dev database.
    """
    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_heading("Extraction Test Standard", level=1)
    doc.add_heading("A.1 Test Clause", level=2)
    doc.add_paragraph(f"This clause is about testing extraction end to end. [{uuid.uuid4().hex}]")
    doc.save(buf)
    return buf.getvalue()


def test_extract_endpoint_builds_real_graph():
    client = TestClient(app)
    files = {"file": ("extraction_test.docx", _sample_docx_bytes(), "application/octet-stream")}

    upload = client.post("/documents", files=files)
    document_id = upload.json()["id"]
    assert client.get(f"/documents/{document_id}").json()["status"] == "ready"

    extract_response = client.post(f"/documents/{document_id}/extract")
    assert extract_response.status_code == 202

    status = client.get(f"/documents/{document_id}").json()
    assert status["graph_status"] == "ready"

    graph = client.get(f"/documents/{document_id}/graph").json()
    assert len(graph["relations"]) == 1
    relation = graph["relations"][0]
    assert relation["source"]["name"] == "Test Control"
    assert relation["target"]["name"] == "Test Risk"
    assert relation["relation_type"] == "APPLIES_TO"


def test_extract_endpoint_404s_for_missing_document():
    client = TestClient(app)
    response = client.post(f"/documents/{uuid.uuid4()}/extract")
    assert response.status_code == 404
