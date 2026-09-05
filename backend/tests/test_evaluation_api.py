"""Live integration tests for Phase 7's evaluation harness.

Deliberately does NOT reuse every other test file's `_mock_external_apis`
convention (faking answer/embedding/judge/classifier LLM calls) — faking the
very thing being measured would make these tests assert only "the harness
calls a mock and gets a mock number back," defeating the feature's entire
purpose. This is a documented, deliberate one-off exception, not a new house
style: every other phase's tests should keep mocking paid APIs as before.

Adds a parallel `_real_llm_available()` skip condition (checks real API keys
are actually configured) alongside the usual `_infra_available()` — a
machine without real Voyage/Cohere/Gemini keys skips cleanly instead of
failing, same discipline `_infra_available()` applies to missing
Docker/infra.
"""

import uuid

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import text

from app.core.config import get_settings
from app.core.db import SessionLocal, engine
from app.main import app
from app.models.evaluation import EvalQuestion
from app.services import graph_store, token_tracking, vector_store
from app.services.query_orchestration import run_query
from app.tasks.celery_app import celery_app


def _infra_available() -> bool:
    try:
        with engine.connect() as conn:
            conn.execute(text("SELECT 1 FROM eval_questions LIMIT 1"))
        vector_store.get_qdrant_client().get_collections()
        driver = graph_store.get_neo4j_driver()
        driver.verify_connectivity()
        driver.close()
        return True
    except Exception:
        return False


def _real_llm_available() -> bool:
    settings = get_settings()
    return bool(settings.gemini_api_key and settings.voyage_api_key and settings.cohere_api_key)


pytestmark = [
    pytest.mark.skipif(
        not _infra_available(),
        reason="requires docker compose up (postgres+redis+qdrant+neo4j) with migrations applied",
    ),
    pytest.mark.skipif(
        not _real_llm_available(),
        reason="requires real GEMINI_API_KEY/VOYAGE_API_KEY/COHERE_API_KEY — "
        "this file deliberately does not mock paid LLM calls",
    ),
]


@pytest.fixture(autouse=True, scope="module")
def _eager_celery():
    celery_app.conf.task_always_eager = True
    celery_app.conf.task_eager_propagates = True
    yield
    celery_app.conf.task_always_eager = False
    celery_app.conf.task_eager_propagates = False


def _seed_standard(nonce: str) -> None:
    """A tiny real standard doc so real retrieval has something to find —
    same nonce-in-clause-text technique test_finding_validation_api.py uses,
    for the same reason (the shared dev/test corpus already holds real
    standards content from earlier phases; a nonce keeps this test's own
    content findable without depending on exactly what else is ingested).
    """
    client = TestClient(app)
    buf_bytes = _sample_docx_bytes(nonce)
    upload = client.post(
        "/documents",
        files={"file": (f"eval_standard_{nonce}.docx", buf_bytes, "application/octet-stream")},
    )
    assert client.get(f"/documents/{upload.json()['id']}").json()["status"] == "ready"


def _sample_docx_bytes(nonce: str) -> bytes:
    import io

    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_heading("Eval Test Standard", level=1)
    doc.add_heading("A.8.24 Use of cryptography", level=2)
    doc.add_paragraph(
        f"Rules for the effective use of cryptography {nonce}, including key "
        f"management {nonce}, shall be defined and implemented."
    )
    doc.save(buf)
    return buf.getvalue()


def _seed_eval_question(nonce: str) -> str:
    db = SessionLocal()
    try:
        question = EvalQuestion(
            question=f"What does the standard say about cryptography {nonce}?",
            use_case_category="audit_evidence_lookup",
            ground_truth_answer=(
                f"Rules for the effective use of cryptography {nonce} shall be defined "
                "and implemented."
            ),
            ground_truth_citations=[{"document_filename": f"eval_standard_{nonce}.docx"}],
            source="llm_drafted",
            human_reviewed=False,
        )
        db.add(question)
        db.commit()
        db.refresh(question)
        return str(question.id)
    finally:
        db.close()


def test_list_and_update_eval_question():
    client = TestClient(app)
    nonce = uuid.uuid4().hex[:8]
    question_id = _seed_eval_question(nonce)

    listed = client.get("/eval/questions", params={"human_reviewed": False}).json()
    assert any(q["id"] == question_id for q in listed)

    response = client.patch(
        f"/eval/questions/{question_id}", json={"human_reviewed": True, "reviewer_name": "tester"}
    )
    assert response.status_code == 200
    assert response.json()["human_reviewed"] is True
    assert response.json()["reviewer_name"] == "tester"


def test_update_eval_question_404s_for_bogus_id():
    client = TestClient(app)
    response = client.patch(
        f"/eval/questions/{uuid.uuid4()}", json={"human_reviewed": True}
    )
    assert response.status_code == 404


def test_create_and_run_eval_run_scores_land_in_valid_range():
    client = TestClient(app)
    nonce = uuid.uuid4().hex[:8]
    _seed_standard(nonce)
    question_id = _seed_eval_question(nonce)

    response = client.post("/eval/runs", json={"question_ids": [question_id]})
    assert response.status_code == 202
    run_id = response.json()["id"]

    run = client.get(f"/eval/runs/{run_id}").json()
    assert run["status"] == "completed"
    assert run["question_count"] == 1
    for field in (
        "avg_faithfulness",
        "avg_answer_relevance",
        "avg_context_precision",
        "avg_context_recall",
    ):
        assert run[field] is None or 0.0 <= run[field] <= 1.0

    results = client.get(f"/eval/runs/{run_id}/results").json()
    assert len(results) == 1
    result = results[0]
    assert result["error_message"] is None
    assert result["generated_answer"]
    assert result["prompt_tokens"] > 0


def test_compare_runs_returns_sane_deltas():
    client = TestClient(app)
    nonce = uuid.uuid4().hex[:8]
    _seed_standard(nonce)
    question_id = _seed_eval_question(nonce)

    run_a_id = client.post("/eval/runs", json={"question_ids": [question_id]}).json()["id"]
    run_b_id = client.post("/eval/runs", json={"question_ids": [question_id]}).json()["id"]

    response = client.get("/eval/runs/compare", params={"a": run_a_id, "b": run_b_id})
    assert response.status_code == 200
    body = response.json()
    assert {d["metric"] for d in body["deltas"]} == {
        "faithfulness",
        "answer_relevance",
        "context_precision",
        "context_recall",
        "latency_ms",
    }


def test_compare_runs_404s_when_either_run_missing():
    client = TestClient(app)
    response = client.get(
        "/eval/runs/compare", params={"a": str(uuid.uuid4()), "b": str(uuid.uuid4())}
    )
    assert response.status_code == 404


def test_token_tracking_captures_nonzero_counts_on_a_real_query():
    # Deliberately a narrow, single-clause factual lookup — the exact shape
    # `query_classifier`'s own system prompt gives as its "vector" example
    # ("what does clause 6.1.2 require?") — not a broad "what is X" question.
    # A broader phrasing risks the real classifier routing to `agent`,
    # which needs the LangGraph Postgres checkpointer; in the full test
    # suite that pool is permanently closed by `test_conversations_api.py`'s
    # own fixture (module-scoped, runs earlier alphabetically), so this test
    # — which calls `run_query` directly with no per-question error
    # isolation, unlike the eval Celery task — must reliably avoid the
    # agent path rather than depend on non-deterministic LLM classification.
    db = SessionLocal()
    try:
        result = run_query("What does clause 4.1 of ISO 27001 require?", db)
    finally:
        db.close()
    assert result.token_usage.prompt_tokens > 0
    # Not asserting completion_tokens > 0 too: an off-topic-classified
    # question's reply is generated in the same classification call, so
    # this only guarantees the classifier's own usage was captured.
    assert token_tracking.current() is result.token_usage
