import io

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import text

from app.core.db import engine
from app.main import app
from app.services import embedding, storage, vector_store
from app.tasks.celery_app import celery_app


def _infra_available() -> bool:
    """Checks Postgres has the migration applied (not just that it's up) and
    that MinIO is reachable. If either fails, this test module is skipped
    rather than failing — it needs the real docker-compose stack, which CI
    doesn't run yet.
    """
    try:
        with engine.connect() as conn:
            conn.execute(text("SELECT 1 FROM documents LIMIT 1"))
        storage.get_minio_client().list_buckets()
        return True
    except Exception:
        return False


pytestmark = pytest.mark.skipif(
    not _infra_available(),
    reason="requires docker compose up (postgres+redis+minio) with migrations applied",
)


@pytest.fixture(autouse=True, scope="module")
def _eager_celery():
    # Runs the parse->chunk->embed chain synchronously in-process, so the
    # test can assert on the final state without a live worker or polling.
    celery_app.conf.task_always_eager = True
    celery_app.conf.task_eager_propagates = True
    yield
    celery_app.conf.task_always_eager = False
    celery_app.conf.task_eager_propagates = False


@pytest.fixture(autouse=True)
def _mock_embed_stage(monkeypatch):
    """This module tests the parse+chunk pipeline, not embeddings — mocking
    Voyage/Qdrant here means the test needs no real API key and no reachable
    Qdrant, while the embed stage still runs (so `document.status` still
    ends up `ready`, matching real pipeline behavior end to end).
    """
    def _fake_embed_texts(texts, input_type):
        return [[0.0] * embedding.EMBEDDING_DIM for _ in texts]

    monkeypatch.setattr(embedding, "embed_texts", _fake_embed_texts)
    monkeypatch.setattr(vector_store, "upsert_chunks", lambda client, chunks, vectors: None)


def _sample_docx_bytes() -> bytes:
    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_heading("Sample Standard", level=1)
    doc.add_heading("A.1 Test Clause", level=2)
    doc.add_paragraph("This is a test clause body.")
    doc.save(buf)
    return buf.getvalue()


def test_upload_document_processes_end_to_end():
    client = TestClient(app)
    data = _sample_docx_bytes()
    docx_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    response = client.post("/documents", files={"file": ("sample.docx", data, docx_type)})
    assert response.status_code == 201
    document_id = response.json()["id"]
    # Enqueued but not yet processed at response time — that's correct even
    # in eager mode, since the chain runs inside apply_async(), which happens
    # after this response body was already built.
    assert response.json()["status"] == "pending"

    status_response = client.get(f"/documents/{document_id}")
    assert status_response.json()["status"] == "ready"

    chunks = client.get(f"/documents/{document_id}/chunks").json()
    assert len(chunks) == 1
    assert chunks[0]["clause_number"] == "A.1"


def test_upload_same_file_twice_is_idempotent():
    client = TestClient(app)
    data = _sample_docx_bytes()

    files = {"file": ("sample.docx", data, "application/octet-stream")}
    first = client.post("/documents", files=files)
    second = client.post("/documents", files=files)

    assert first.json()["id"] == second.json()["id"]


def test_rerunning_chunk_stage_does_not_duplicate_chunks():
    """Real bug found live: resubmitting the parse->chunk->embed chain to
    retry a document that failed at the embed stage (chunking having
    already succeeded) re-ran chunk_document_task, which didn't clear its
    prior rows first — doubling every chunk. This directly exercises the
    fix: rerunning the chunk stage in isolation must be a no-op on chunk
    count, not additive.
    """
    from app.tasks.ingestion import chunk_document_task, parse_document_task

    client = TestClient(app)
    data = _sample_docx_bytes()
    docx_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response = client.post("/documents", files={"file": ("rerun.docx", data, docx_type)})
    document_id = response.json()["id"]
    assert client.get(f"/documents/{document_id}").json()["status"] == "ready"

    first_count = len(client.get(f"/documents/{document_id}/chunks").json())
    assert first_count == 1

    parsed_tree = parse_document_task(document_id)
    chunk_document_task(parsed_tree, document_id)

    second_count = len(client.get(f"/documents/{document_id}/chunks").json())
    assert second_count == first_count
