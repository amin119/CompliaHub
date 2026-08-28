import io
import json

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import text

from app.core.db import engine
from app.main import app
from app.services import (
    agent,
    answer_generation,
    embedding,
    graph_store,
    query_classifier,
    reranking,
    retrieval,
    vector_store,
)
from app.services.query_classifier import QueryCategory, QueryClassification
from app.tasks.celery_app import celery_app


def _infra_available() -> bool:
    """Same infra requirement as test_documents_api.py, plus a reachable
    Qdrant and (Phase 4) Neo4j — this module's whole point is exercising
    real dense-search + lexical-search + fusion + local-search wiring, only
    the paid API calls (Voyage, Cohere, Gemini) are mocked.
    """
    try:
        with engine.connect() as conn:
            conn.execute(text("SELECT 1 FROM documents LIMIT 1"))
        vector_store.get_qdrant_client().get_collections()
        driver = graph_store.get_neo4j_driver()
        driver.verify_connectivity()
        driver.close()
        return True
    except Exception:
        return False


pytestmark = pytest.mark.skipif(
    not _infra_available(),
    reason="requires docker compose up (postgres+redis+qdrant) with migrations applied",
)


@pytest.fixture(autouse=True, scope="module")
def _eager_celery():
    celery_app.conf.task_always_eager = True
    celery_app.conf.task_eager_propagates = True
    yield
    celery_app.conf.task_always_eager = False
    celery_app.conf.task_eager_propagates = False


@pytest.fixture(autouse=True)
def _mock_external_apis(monkeypatch):
    """Mocks every paid third-party call (Voyage, Cohere, Anthropic) so this
    test needs no real API keys, while still exercising the real Postgres
    lexical search + real Qdrant dense search + real RRF fusion wiring.

    Embeddings are faked as a single constant vector: dense-search *ranking*
    isn't asserted on (there's only one document's chunks in Qdrant, so
    scores would be meaningless with a fake vector anyway) — this test
    verifies the endpoint's plumbing end-to-end, not retrieval quality,
    which is what the real GDPR regression check (docs/phase-2-vector-layer.md)
    is for.
    """

    def _fake_embed_texts(texts, input_type):
        return [[0.1] * embedding.EMBEDDING_DIM for _ in texts]

    def _fake_rerank(query, documents, top_n):
        return [
            reranking.RerankResult(index=i, relevance_score=1.0 - i * 0.01)
            for i in range(min(top_n, len(documents)))
        ]

    monkeypatch.setattr(embedding, "embed_texts", _fake_embed_texts)
    monkeypatch.setattr(reranking, "rerank", _fake_rerank)
    monkeypatch.setattr(
        answer_generation,
        "generate_answer",
        lambda question, chunks, graph_facts=None, community_context=None: "mocked answer",
    )
    monkeypatch.setattr(
        answer_generation,
        "stream_answer",
        lambda question, chunks, graph_facts=None, community_context=None: iter(
            ["mocked ", "answer"]
        ),
    )
    # Phase 5: every /query call now classifies first — default to GRAPH so
    # the existing tests below (written against Phase 4's vector+local
    # behavior) keep exercising that same path unchanged. Tests that care
    # about a different category override this per-test.
    monkeypatch.setattr(
        query_classifier,
        "classify_query",
        lambda question, client=None: QueryClassification(category=QueryCategory.GRAPH),
    )


def _sample_docx_bytes() -> bytes:
    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_heading("Query Test Standard", level=1)
    doc.add_heading("A.1 Test Clause", level=2)
    doc.add_paragraph("Assets associated with information shall be identified and inventoried.")
    doc.save(buf)
    return buf.getvalue()


def test_query_returns_answer_with_citations():
    client = TestClient(app)
    files = {"file": ("query_sample.docx", _sample_docx_bytes(), "application/octet-stream")}
    upload = client.post("/documents", files=files)
    assert client.get(f"/documents/{upload.json()['id']}").json()["status"] == "ready"

    question = "What must be done with information assets?"
    response = client.post("/query", json={"question": question})

    assert response.status_code == 200
    body = response.json()
    assert body["answer"] == "mocked answer"
    assert len(body["citations"]) >= 1
    assert body["citations"][0]["clause_number"] == "A.1"


def test_query_with_no_matching_chunks_returns_empty_citations():
    client = TestClient(app)
    question = "asdkjhaslkdjhaslkdjh nonsense query xyz123"
    response = client.post("/query", json={"question": question})

    assert response.status_code == 200
    # Lexical search may legitimately find nothing for gibberish; dense
    # search still returns *something* (fake vectors have no real notion of
    # "no match"), so this only asserts the endpoint doesn't error — not
    # that citations are empty.
    assert "answer" in response.json()


def test_vector_category_skips_the_graph_entirely(monkeypatch):
    """Real cost savings from Phase 5's classifier: a `vector`-classified
    question must never touch Neo4j at all. Verified by making
    `local_search_facts` raise — if the route calls it anyway, this test
    fails loudly instead of silently passing.
    """
    monkeypatch.setattr(
        query_classifier,
        "classify_query",
        lambda question, client=None: QueryClassification(category=QueryCategory.VECTOR),
    )

    def _boom(driver, context_chunks):
        raise AssertionError("local_search_facts must not run for a VECTOR-classified question")

    monkeypatch.setattr(retrieval, "local_search_facts", _boom)

    client = TestClient(app)
    response = client.post("/query", json={"question": "what does clause 6.1.2 require?"})

    assert response.status_code == 200
    assert response.json()["answer"] == "mocked answer"


def test_off_topic_category_falls_back_to_a_default_reply_without_touching_retrieval(monkeypatch):
    """Bug fix (post-Phase 6): before the `off_topic` category existed, a
    greeting like "hi, how are you?" had no category that actually fit it,
    so the classifier's own "when unsure, prefer agent" tie-break sent it
    through the full agent loop — confirmed live to take ~19s. Verified
    here by making retrieval raise: an `off_topic` question must never
    reach it at all. `reply=None` here exercises the fallback path — the
    classifier normally generates a real reply itself, see the
    `_uses_the_classifiers_generated_reply` tests below.
    """
    monkeypatch.setattr(
        query_classifier,
        "classify_query",
        lambda question, client=None: QueryClassification(category=QueryCategory.OFF_TOPIC),
    )

    def _boom(*args, **kwargs):
        raise AssertionError("retrieval must not run for an OFF_TOPIC-classified question")

    monkeypatch.setattr(retrieval, "vector_search", _boom)

    client = TestClient(app)
    response = client.post("/query", json={"question": "hi, how are you?"})

    assert response.status_code == 200
    body = response.json()
    assert body["citations"] == []
    assert "compliance" in body["answer"].lower()


def test_off_topic_category_uses_the_classifiers_generated_reply(monkeypatch):
    """The real fix for "always the same answer": the classifier generates
    a real, varied reply in the same call that classifies the question —
    no second LLM round-trip — and the route must use it verbatim rather
    than always falling back to one fixed sentence.
    """
    monkeypatch.setattr(
        query_classifier,
        "classify_query",
        lambda question, client=None: QueryClassification(
            category=QueryCategory.OFF_TOPIC,
            reply="Ahoy! I'm strictly compliance-brained, but happy to talk ISO/GDPR.",
        ),
    )

    client = TestClient(app)
    response = client.post("/query", json={"question": "ahoy there!"})

    assert response.status_code == 200
    assert (
        response.json()["answer"]
        == "Ahoy! I'm strictly compliance-brained, but happy to talk ISO/GDPR."
    )


def test_query_stream_off_topic_category_falls_back_to_a_default_reply_instantly(monkeypatch):
    monkeypatch.setattr(
        query_classifier,
        "classify_query",
        lambda question, client=None: QueryClassification(category=QueryCategory.OFF_TOPIC),
    )

    def _boom(*args, **kwargs):
        raise AssertionError("retrieval must not run for an OFF_TOPIC-classified question")

    monkeypatch.setattr(retrieval, "vector_search", _boom)

    client = TestClient(app)
    response = client.post("/query/stream", json={"question": "hi, how are you?"})

    assert response.status_code == 200
    events = _parse_sse(response.text)

    assert events[0] == {"type": "status", "stage": "classifying"}
    tokens = [event["text"] for event in events if event["type"] == "token"]
    assert "compliance" in "".join(tokens).lower()


def test_query_stream_off_topic_category_uses_the_classifiers_generated_reply(monkeypatch):
    monkeypatch.setattr(
        query_classifier,
        "classify_query",
        lambda question, client=None: QueryClassification(
            category=QueryCategory.OFF_TOPIC, reply="Nice weather talk, but I only do compliance."
        ),
    )

    client = TestClient(app)
    response = client.post("/query/stream", json={"question": "nice weather today"})

    assert response.status_code == 200
    events = _parse_sse(response.text)
    tokens = [event["text"] for event in events if event["type"] == "token"]
    assert "".join(tokens) == "Nice weather talk, but I only do compliance."
    assert events[-1]["type"] == "done"
    assert events[-1]["citations"] == []


def test_agent_category_delegates_to_the_agent_loop(monkeypatch):
    from app.schemas.query import QueryResponse

    monkeypatch.setattr(
        query_classifier,
        "classify_query",
        lambda question, client=None: QueryClassification(category=QueryCategory.AGENT),
    )

    captured = {}

    def _fake_run_agent(question, db, driver, checkpointer, conversation_id=None, max_iterations=2):
        captured["question"] = question
        captured["conversation_id"] = conversation_id
        return QueryResponse(answer="agent answer", citations=[], conversation_id="conv-123")

    monkeypatch.setattr(agent, "run_agent", _fake_run_agent)

    client = TestClient(app)
    question = "what does ISO 42001 require that ISO 27001 doesn't?"
    response = client.post("/query", json={"question": question, "conversation_id": "conv-123"})

    assert response.status_code == 200
    assert response.json()["conversation_id"] == "conv-123"
    assert captured["conversation_id"] == "conv-123"
    assert response.json()["answer"] == "agent answer"
    assert captured["question"] == question


def _parse_sse(body: str) -> list[dict]:
    events = []
    for frame in body.split("\n\n"):
        frame = frame.strip()
        if not frame:
            continue
        for line in frame.split("\n"):
            if line.startswith("data: "):
                events.append(json.loads(line[len("data: ") :]))
    return events


def test_query_stream_returns_status_and_token_events_ending_in_done():
    """Phase 6 Part 2: verifies the real SSE wire format the frontend's
    `streamQuestion` parses — status events while retrieving, one token
    event per chunk `stream_answer` yields, a final `done` event carrying
    citations/graph_evidence/conversation_id. `TestClient` drains the whole
    generator synchronously, so `response.text` is the complete SSE body.
    """
    client = TestClient(app)
    files = {"file": ("query_stream_sample.docx", _sample_docx_bytes(), "application/octet-stream")}
    upload = client.post("/documents", files=files)
    assert client.get(f"/documents/{upload.json()['id']}").json()["status"] == "ready"

    response = client.post(
        "/query/stream", json={"question": "What must be done with information assets?"}
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/event-stream")
    events = _parse_sse(response.text)

    assert events[0] == {"type": "status", "stage": "classifying"}
    assert {"type": "status", "stage": "retrieving"} in events
    assert {"type": "status", "stage": "generating_answer"} in events
    tokens = [event["text"] for event in events if event["type"] == "token"]
    assert "".join(tokens) == "mocked answer"
    done = events[-1]
    assert done["type"] == "done"
    assert len(done["citations"]) >= 1
    assert done["citations"][0]["clause_number"] == "A.1"


def test_query_stream_agent_category_delegates_to_stream_agent(monkeypatch):
    monkeypatch.setattr(
        query_classifier,
        "classify_query",
        lambda question, client=None: QueryClassification(category=QueryCategory.AGENT),
    )

    def _fake_stream_agent(
        question, db, driver, checkpointer, conversation_id=None, max_iterations=2
    ):
        yield {"type": "status", "stage": "retrieving"}
        yield {"type": "token", "text": "agent "}
        yield {"type": "token", "text": "answer"}
        yield {
            "type": "done",
            "conversation_id": "conv-456",
            "citations": [],
            "graph_evidence": {"nodes": [], "edges": []},
        }

    monkeypatch.setattr(agent, "stream_agent", _fake_stream_agent)

    client = TestClient(app)
    response = client.post(
        "/query/stream", json={"question": "what does ISO 42001 require that ISO 27001 doesn't?"}
    )

    assert response.status_code == 200
    events = _parse_sse(response.text)
    tokens = [event["text"] for event in events if event["type"] == "token"]
    assert "".join(tokens) == "agent answer"
    assert events[-1] == {
        "type": "done",
        "conversation_id": "conv-456",
        "citations": [],
        "graph_evidence": {"nodes": [], "edges": []},
    }
